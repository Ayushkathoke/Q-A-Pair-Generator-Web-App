import streamlit as st
from utl import goto
import base64
import tempfile
import pdfplumber
from conf import config
from fpdf import FPDF
from docx import Document
from llm import ChatLL

chat = ChatLL()


def generate_pdf(qna_text):
    # Create a PDF instance
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Set font for PDF
    pdf.set_font("Arial", size=12)

    # Add the Q&A content to the PDF
    pdf.multi_cell(0, 10, qna_text)

    # Save to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf_file:
        pdf.output(temp_pdf_file.name)
        return temp_pdf_file.name


def generate_docx(qna_text):
    # Create a DOCX instance
    doc = Document()

    # Add title and Q&A content
    doc.add_heading("Generated Q&A", 0)
    doc.add_paragraph(qna_text)

    # Save to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx_file:
        doc.save(temp_docx_file.name)
        return temp_docx_file.name


def HomePage():
    st.set_page_config(
        page_title="Question Answering Generator", layout="wide")
    st.title('Question Answering Generator:books:')
    st.markdown('<style>h1{color: orange; text-align: center;}</style>', unsafe_allow_html=True)
    st.subheader('Built for Professionals,Teachers, Students')
    st.markdown('<style>h3{color: pink;  text-align: center;}</style>', unsafe_allow_html=True)
    # Sidebar settings
    st.sidebar.header("Settings")
    uploaded_file = st.sidebar.file_uploader("Upload a PDF file 🏠", type=["pdf"])
    difficulty_level = st.sidebar.selectbox(
        "Select Difficulty Level", ["Easy", "Moderate", "Hard"])
    num_questions = st.sidebar.number_input(
        "Number of Questions", min_value=1, max_value=100, value=5)
    question_type = st.sidebar.selectbox("Select Question Type", [
                                         "Descriptive", "One-Short", "MCQ", "True/False", "Yes/No", "Fill in the Blanks"])

    # Navigation button to go back to Home
    st.sidebar.button("Go to Home", icon="🏠", on_click=goto, args=("home",))
    
    generate_button = st.sidebar.button("Generate Questions", icon="🏠")

    text = ""  # Initialize text variable outside the file check

    if uploaded_file:
        # Save the uploaded file to a temporary location
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(uploaded_file.read())
            file_path = temp_file.name

        if file_path:
            # Handle PDF file and extract text
            with uploaded_file as pdf_file:
                with pdfplumber.open(pdf_file) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text()

            # Display the PDF content using base64 encoding
            with open(file_path, "rb") as f:
                base64_pdf = base64.b64encode(f.read()).decode('utf-8')
            pdf_display = f'''
            <div style="display: flex; justify-content: center;">
                <embed src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf">
            </div>
            '''
            st.markdown(pdf_display, unsafe_allow_html=True)

            # Display extracted text (for debugging purposes)
            st.text_area("Extracted PDF Text", text, height=200)

    # Logic when Generate Questions button is clicked
    if generate_button:
        if not uploaded_file:
            st.warning("Please upload a PDF file.", icon="😅")
        elif not text or text == "":
            st.warning(
                "Failed to extract text from the uploaded PDF.", icon="😅")
        elif not difficulty_level:
            st.warning("Please select a difficulty level.", icon="😅")
        elif not question_type:
            st.warning("Please select a question type.", icon="😅")
        elif not num_questions:
            st.warning("Please specify the number of questions.", icon="😅")
        else:
            _text = chat.runner(int(num_questions), context=text,
                                difficultyLevel=difficulty_level, questionType=question_type)
            st.markdown(_text, unsafe_allow_html=True)
            download_option = st.radio("Download as", ("PDF", "DOCX"))

            if download_option == "PDF":
                pdf_file_path = generate_pdf(_text)
                with open(pdf_file_path, "rb") as pdf_file:
                    pdf_data = pdf_file.read()
                    st.download_button(
                        label="Download PDF",
                        data=pdf_data,
                        file_name="generated_questions.pdf",
                        mime="application/pdf"
                    )

            elif download_option == "DOCX":
                docx_file_path = generate_docx(_text)
                with open(docx_file_path, "rb") as docx_file:
                    docx_data = docx_file.read()
                    st.download_button(
                        label="Download DOCX",
                        data=docx_data,
                        file_name="generated_questions.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
